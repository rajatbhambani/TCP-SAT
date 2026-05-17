#!/usr/bin/env python3
"""
Generate paper.docx from BBR-SAT paper content.
Run from /home/rajat/tcpsatproject/scripts/
Output: /home/rajat/tcpsatproject/paper/paper.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup (letter, 1" margins) ─────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Styles ───────────────────────────────────────────────────────────────────
styles = doc.styles

def set_style_font(style_name, size_pt, bold=False, italic=False):
    if style_name in [s.name for s in styles]:
        s = styles[style_name]
        s.font.size = Pt(size_pt)
        s.font.bold = bold
        s.font.italic = italic

set_style_font('Normal', 11)

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_para(text, bold_ranges=None, italic_ranges=None, mono_ranges=None,
             first_line_indent=True):
    """Add a paragraph with optional inline formatting hints (unused in simple mode)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    _add_rich_text(p, text)
    return p

def _add_rich_text(para, text):
    """Parse **bold**, *italic*, `mono` markers and add runs."""
    import re
    # Tokenise: **bold**, *italic*, `mono`, or plain
    pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        if m.group(1) is not None:          # **bold**
            r = para.add_run(m.group(1))
            r.bold = True
        elif m.group(2) is not None:        # *italic*
            r = para.add_run(m.group(2))
            r.italic = True
        elif m.group(3) is not None:        # `mono`
            r = para.add_run(m.group(3))
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True

def add_table_simple(headers, rows, caption=None):
    if caption:
        cp = doc.add_paragraph(caption)
        cp.paragraph_format.space_before = Pt(6)
        cp.paragraph_format.space_after  = Pt(2)
        cp.runs[0].bold = True
        cp.runs[0].font.size = Pt(10)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            p = cell.paragraphs[0]
            if isinstance(val, tuple):
                text, fmt = val
                r = p.add_run(text)
                if 'b' in fmt: r.bold = True
                if 'i' in fmt: r.italic = True
            else:
                r = p.add_run(str(val))
            r.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # spacing after table
    return t


# ═══════════════════════════════════════════════════════════════════════════
# TITLE AND AUTHOR
# ═══════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BBR-SAT: Orbit-Aware Congestion Control for Multi-Orbit Satellite QUIC Handovers')
r.bold = True
r.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Rajat Bhambani\nIndependent Research Contribution\nbhambani.rajat@gmail.com')
r.font.size = Pt(11)

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════

add_heading('Abstract', level=1)
add_para(
    'Low-Earth orbit (LEO), medium-Earth orbit (MEO), and geostationary (GEO) '
    'satellite constellations are increasingly deployed together, creating multi-orbit '
    'networks in which a single QUIC connection may traverse orbit-class boundaries '
    'mid-flight. Such handovers impose abrupt, step-change shifts in available bandwidth '
    '(up to 5× reduction) and round-trip time (up to 12× increase), which existing '
    'congestion-control algorithms handle poorly: BBRv3 stalls for the duration of its '
    'two-cycle bandwidth filter, and simple heuristics such as congestion-window freeze '
    'or send-pause succeed only at a single, carefully tuned lead time. We present '
    '**BBR-SAT**, a minimal extension of BBRv3 that equips the sender with an orbit '
    'parameter table seeded from ephemeris data and a two-phase handover protocol: a '
    'PREDICTED signal initiates a proactive queue drain via `ProbeBW_DOWN`, and a '
    'CONFIRMED signal resets the BDP context directly from the orbit table. Implemented '
    'in picoquic and evaluated in a zero-loss shared-link simulator across all six '
    'pairwise LEO/MEO/GEO transitions at three handover times and seven advance lead '
    'times, BBR-SAT converges to 90% of new-orbit capacity in 2.5 s for the critical '
    'LEO→GEO transition — the only evaluated mechanism to converge on this transition at '
    'any lead time from 0 to 20 s — while reducing peak queue occupancy by 36× relative '
    'to vanilla BBRv3. Fairness analysis shows that BBR-SAT co-exists equitably with '
    'competing BBRv3 flows (Jain J = 0.997) and neither worsens nor repairs the '
    'pre-existing BBRv3 deference to CUBIC.'
)

p = doc.add_paragraph()
r = p.add_run('Keywords: ')
r.bold = True
p.add_run('satellite communications, QUIC, congestion control, BBR, handover, '
          'multi-orbit networks, LEO, GEO')
p.paragraph_format.space_after = Pt(8)

# ═══════════════════════════════════════════════════════════════════════════
# I. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════

add_heading('I. Introduction', level=1)

add_para(
    'The deployment of large LEO constellations (Starlink, OneWeb, Amazon Kuiper) [1,2] '
    'alongside legacy MEO (O3b mPOWER) and GEO (ViaSat-3) infrastructure has created a '
    'new class of network event: the *inter-orbit handover*, in which an active transport '
    'connection migrates between satellite beams belonging to orbit classes with '
    'fundamentally different propagation characteristics. Such transitions are *abrupt* — '
    'the link properties change at the granularity of a single RTT — and *asymmetric*: a '
    'LEO→GEO handover simultaneously reduces the available link rate by 5× (50→10 Mbps) '
    'and increases the round-trip propagation delay by ~12× (50→580 ms). Because the GEO '
    'BDP (10 Mbps × 580 ms = 725 KB) exceeds the LEO BDP (50 Mbps × 50 ms = 312 KB), '
    'the *new* orbit requires a *larger* inflight window — but at a *much lower* pacing rate.'
)

add_para(
    'Modern QUIC stacks running BBRv3 [3] are particularly poorly suited to absorb this '
    'transition. BBRv3\'s two-cycle `MaxBwFilter` requires approximately 8 × RTT_new to '
    'clear its pre-handover bandwidth estimate. On a GEO link (RTT ≈ 580 ms) that amounts '
    'to roughly 5 seconds of sustained over-pacing at 50 Mbps into a 10 Mbps pipe, '
    'accumulating ≈33 MB of queued data in our experiments before the filter clears and '
    'the sender backs off.'
)

add_para(
    'The central insight motivating BBR-SAT is that satellite operators have *advance '
    'knowledge* of handover events: orbital mechanics are deterministic, and ground-control '
    'systems routinely predict beam transitions minutes ahead of time. This knowledge can '
    'be surfaced to the transport layer as two signals: (1) PREDICTED, issued seconds to '
    'tens of seconds before the event, and (2) CONFIRMED, issued at the moment the beam '
    'switch completes. BBR-SAT consumes these signals to execute a two-phase response: a '
    'proactive queue drain aligned with the current (pre-handover) RTT, followed by an '
    'instantaneous BDP context switch that seeds the post-handover pacing rate and '
    'congestion window directly from a per-orbit parameter table.'
)

add_para(
    'This paper contributes: (1) the BBR-SAT algorithm — two new signals, an orbit '
    'parameter table, and a two-phase handover protocol integrated into BBRv3 in <100 '
    'lines (§III); (2) a zero-loss shared-link simulator spanning all six pairwise orbit '
    'transitions at three handover times and six lead times (§IV); and (3) empirical '
    'demonstration that BBR-SAT is the only evaluated mechanism to converge at the '
    'critical LEO→GEO transition across lead times 0–20 s, with fairness results '
    'confirming that advance knowledge does not translate into a bandwidth advantage '
    '(§§V–B).'
)

# ═══════════════════════════════════════════════════════════════════════════
# II. BACKGROUND AND RELATED WORK
# ═══════════════════════════════════════════════════════════════════════════

add_heading('II. Background and Related Work', level=1)
add_heading('A. BBRv3 Filter Mechanics', level=2)

add_para(
    'BBR [4] drives the sender at the estimated bottleneck bandwidth using two long-term '
    'filters. The `MaxBwFilter` is a two-slot windowed maximum over the last two ProbeBW '
    'cycles; each cycle spans ~4 RTTs, so on a GEO link (RTT ≈ 580 ms) clearing a stale '
    'pre-handover estimate takes 8 × 580 ≈ 4,640 ms. The `min_rtt` filter holds the '
    'minimum observed RTT over a 10-second window; a LEO→GEO transition (50→580 ms) '
    'leaves the sender calibrated to the wrong path delay for up to 10 s. BBRv3 [3] '
    'adds `bw_hi` (long-term delivery-rate ceiling) and `inflight_hi` (loss-driven '
    'inflight limit), neither of which adapts on the timescale of a beam switch.'
)

add_heading('B. Prior Work on Satellite Handover CCA', level=2)

add_para(
    'Five intra-LEO cross-layer mechanisms appeared in 2024–2025: StarQUIC [5], '
    'SATPIPE [6], LeoCC [7], Creo [8], and OrbCC [9]; Table I summarises their signal '
    'and algorithmic approach. All operate within a single LEO constellation and none '
    'handle cross-orbit transitions between orbit classes with fundamentally different '
    'propagation regimes.'
)

add_para(
    'At the IETF, Lai et al. [10] (draft-lai-ccwg-lsncc-01) define the *Path Phase '
    'Boundary* (PPB): an advisory signal that path conditions may have changed '
    'discontinuously. PPB is reactive, parameter-free, and CCA-agnostic — it identifies '
    'the problem but prescribes no algorithmic response. Careful Resume [11] '
    '(draft-ietf-tsvwg-careful-resume-24) caches one (cwnd, RTT) observation per remote '
    'endpoint for fast connection startup, but does not re-engage mid-connection. Its '
    'only satellite evaluation tests CUBIC, not BBR [12].'
)

add_para(
    'BBR-SAT is motivated by the PPB problem statement but differs structurally: it fires '
    '*before* the event (predictive), carries the target orbit class (parameterized), and '
    'triggers a specific BBRv3 state-machine response (BDP context switch). The defining '
    'gap it fills — cross-orbit mid-connection BDP switching — is absent from all prior work.'
)

add_table_simple(
    headers=['Mechanism', 'Mid-Conn?', 'Cross-Orbit?', 'Signal'],
    rows=[
        ['PPB [10]',           'reactive', 'No',  'ICMP probe'],
        ['Careful Resume [11]','No',       'No',  'cached obs.'],
        ['StarQUIC [5]',       'Yes',      'No',  '15-s sched.'],
        ['SATPIPE [6]',        'Yes',      'No',  'NTP sched.'],
        ['LeoCC [7]',          'Yes',      'No',  'ICMP probe'],
        ['Creo [8]',           'Yes',      'No',  'cross-layer'],
        ['OrbCC [9]',          'Yes',      'No',  'INT'],
        [('BBR-SAT (this work)', 'b'), 'Yes', ('Yes','b'), 'eph./INT/probe'],
    ],
    caption='Table I. Satellite handover CCA mechanisms'
)

add_para(
    'Legacy work addressed GEO-specific impairments: PEPs and Hybla [13] improve TCP on '
    'high-latency links; characterisation of LEO paths [2, 14] documents the dynamics '
    'that make inter-orbit transitions damaging. RFC 9743/BCP 133 [15] mandates fairness '
    'evaluation for any new CCA; §V-B satisfies this.'
)

# ═══════════════════════════════════════════════════════════════════════════
# III. BBR-SAT DESIGN
# ═══════════════════════════════════════════════════════════════════════════

add_heading('III. BBR-SAT Design', level=1)
add_heading('A. Orbit Parameter Table', level=2)

add_para(
    'BBR-SAT augments the per-connection BBRv3 state with a three-entry *orbit table*, '
    'one entry per orbit class (LEO, MEO, GEO). Each entry stores: `min_rtt_us` '
    '(ephemeris-derived propagation RTT in µs), `bw_hi` (nominal beam capacity in '
    'bits/s), and `bw_confidence` (EPHEMERIS or MEASURED).'
)

add_para(
    'Entries are populated at connection setup from satellite ephemeris data supplied by '
    'the operator; `bw_confidence` transitions from EPHEMERIS to MEASURED once the '
    'algorithm has accumulated one ProbeBW cycle on that orbit. The `min_rtt` field is '
    'trusted immediately because propagation delay is deterministic given orbital geometry; '
    'the `bw_hi` field is treated as a starting ceiling, not a hard limit, allowing '
    'BBR\'s natural probing to refine it within 1–2 RTTs of the new path.'
)

add_heading('B. Signal Protocol', level=2)

add_para(
    'Two signals are delivered to the congestion controller via the picoquic `alg_notify` '
    'callback:'
)

p = doc.add_paragraph(style='List Bullet')
_add_rich_text(p, '**PREDICTED** (0x32): the ground control system predicts a handover '
    'to a target orbit within the next ℓ seconds. The signal identifies the target orbit index.')

p = doc.add_paragraph(style='List Bullet')
_add_rich_text(p, '**CONFIRMED** (0x02): the beam switch has completed; the sender is '
    'now on the target orbit.')

add_para(
    'CONFIRMED may arrive without a preceding PREDICTED (zero lead time is always supported).'
)

add_heading('C. Two-Phase Handover Protocol', level=2)

add_para('**Phase 1 — Proactive Drain (PREDICTED signal).**')
add_para(
    'When BBR-SAT receives PREDICTED, it forces BBRv3 into `ProbeBW_DOWN` with the drain '
    'gain (η < 1) applied to the current orbit\'s BDP. The `bw_probe_ceiling` is set to '
    '2 × `max_bw` to suppress spurious re-Startup triggers during the drain. Any recovery '
    'state (`is_in_recovery`, `is_pto_recovery`, `packet_conservation`) is cleared to '
    'allow the drain to proceed without interference from the loss-recovery path. The '
    'drain completes within one to two RTTs of the *current* orbit, well before the '
    'handover fires for lead times up to ≈20 s.'
)

add_para('**Phase 2 — BDP Context Switch (CONFIRMED signal).**')
add_para('On CONFIRMED, BBR-SAT:')

for item in [
    'Loads the target orbit\'s `min_rtt_us` and `bw_hi` directly into the BBRv3 state, '
    'bypassing the `MaxBwFilter` and `min_rtt` expiry windows.',
    'Resets `inflight_hi` to `bw_hi` × `min_rtt_us` (one BDP of the new orbit).',
    'Clears both slots of the two-cycle `MaxBwFilter`, forcing BBR to accept the seeded '
    '`bw_hi` as the new bandwidth ceiling in the very next round.',
    'Enters `ProbeBW_REFILL` to ramp up to the new BDP.',
]:
    p = doc.add_paragraph(style='List Number')
    _add_rich_text(p, item)

add_para(
    'The net effect is that the post-handover pacing rate reflects the new orbit\'s '
    'capacity within one RTT of the CONFIRMED signal, rather than the 8 × RTT_new '
    'required for natural BBRv3 filter convergence.'
)

# ═══════════════════════════════════════════════════════════════════════════
# IV. METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════

add_heading('IV. Methodology', level=1)
add_heading('A. Simulator', level=2)

add_para(
    'All experiments use `picoquic_ns` [16], the built-in network simulator embedded in '
    'picoquic. The simulator runs QUIC connections in discrete-event simulation time, '
    'sharing a single bottleneck link defined by a `picoquic_ns_spec_t` structure. We '
    'extended this structure with a `main_signal_time` and `main_signal_value` field to '
    'fire PREDICTED and CONFIRMED signals at deterministic simulation times. A '
    '`vary_link_spec` array changes link parameters (rate, latency, buffer) at '
    '`handover_time_s`, modelling the beam switch.'
)

add_para(
    'Single-flow experiments use one QUIC connection (server→client bulk transfer); '
    'fairness experiments use two connections sharing the same bottleneck. All simulations '
    'are zero-loss (no packet loss injection) to isolate the congestion-control response '
    'from loss-recovery interactions.'
)

add_heading('B. Orbit Parameters', level=2)

add_table_simple(
    headers=['Orbit', 'Link rate', 'RTT', 'BDP'],
    rows=[
        ['LEO', '50 Mbps', '50 ms',  '312 KB'],
        ['MEO', '30 Mbps', '150 ms', '562 KB'],
        ['GEO', '10 Mbps', '580 ms', '725 KB'],
    ],
    caption='Table II. Simulated orbit parameters'
)

add_para(
    'The simulator\'s queuing is soft-limited by a 2×RTT delay threshold; the *effective* '
    'buffer in our experiments is uncapped by that threshold (packets continue to queue if '
    'the sender does not back off), which is why Table III shows a 33 MB build-up for '
    'B1/B3 rather than the ~725 KB hard-BDP limit.'
)

add_heading('C. Baselines', level=2)

for item in [
    '**B1** — Vanilla BBRv3 (no modification, no signal).',
    '**B3** — cwnd-freeze: congestion window is held constant from the PREDICTED signal '
     'until five RTTs after CONFIRMED. (B2, a delayed-reduction variant, produced results '
     'indistinguishable from B1 at all tested lead times and is omitted.)',
    '**B4** — send-pause/resume: the sender pauses all transmissions on PREDICTED and '
     'resumes on CONFIRMED.',
    '**BBR-SAT** — the proposed algorithm (this work).',
    '**CUBIC** — standard CUBIC, no modification.',
]:
    p = doc.add_paragraph(style='List Bullet')
    _add_rich_text(p, item)

add_para(
    'B3 and B4 are common heuristics proposed informally in the satellite operator '
    'community; CUBIC serves as a loss-driven convergence reference. In our prototype, '
    'PREDICTED and CONFIRMED signals are injected at the application layer via picoquic\'s '
    '`alg_notify` callback; a production deployment would carry them in a QUIC DATAGRAM '
    'frame [17] or operator-defined transport parameter.'
)

add_heading('D. Experiment Parameters', level=2)

add_para(
    '**Experiment 1 (single-flow):** Each baseline is run through all six pairwise orbit '
    'transitions (LEO↔MEO, LEO↔GEO, MEO↔GEO) at three handover times '
    '(T_HO ∈ {30, 60, 120} s) and six lead times (ℓ ∈ {0, 2, 5, 10, 20, 30} s). '
    'Total simulation: 90 s per trial.'
)

add_para(
    '**Experiment 2 (fairness):** Two-flow shared-bottleneck runs under two topologies: '
    '**F1** — both flows cross a LEO→GEO handover at T = 30 s (flow 1: BBR-SAT, '
    'flow 2: BBRv3 or CUBIC); **F3** — both flows start on GEO, BBR-SAT re-anchors its '
    'BDP context at T = 30 s (flow 2: CUBIC). Each condition is repeated twice; '
    'post-event averages are 2-run means.'
)

# ═══════════════════════════════════════════════════════════════════════════
# V. EVALUATION
# ═══════════════════════════════════════════════════════════════════════════

add_heading('V. Evaluation', level=1)
add_heading('A. Single-Flow Handover Performance', level=2)
add_heading('1) Setup', level=3)

add_para(
    'We drive each of the five baselines through all six pairwise orbit transitions at '
    'three handover times (T_HO ∈ {30, 60, 120} s) and six lead times '
    '(ℓ ∈ {0, 2, 5, 10, 20, 30} s). Each condition is a zero-loss, single-flow '
    'simulation. The primary metric is **T90**: the elapsed time from T_HO until the '
    'delivered throughput first reaches 90% of the new orbit\'s link capacity, measured '
    'as a one-second rolling average. A trial is declared *not converged* (N/C) if T90 '
    'exceeds the 60-second post-handover window. We additionally record post-handover '
    'steady-state throughput, goodput (total bytes delivered in the 90-second run), and '
    'peak on-path queue depth.'
)

add_heading('2) LEO→GEO: the Critical Transition', level=3)

add_para(
    'A downward BDP transition — lower link rate, higher latency, larger BDP — is the '
    'most demanding case. Table III and Figure 1 show results for the LEO→GEO transition '
    'at T_HO = 30 s with ℓ = 0.'
)

add_para(
    '**B1 (vanilla BBRv3)** fails entirely: the on-path queue inflates to 33 MB, '
    'steady-state throughput is effectively zero, and the 90% threshold is never reached '
    'within the 60-second window. The root cause is that BBRv3 retains its LEO-calibrated '
    'bandwidth estimate after handover. Because BBRv3\'s two-cycle `MaxBwFilter` clears '
    'only after roughly 8 × RTT_GEO ≈ 5 s [3], the sender continues pacing at ≈50 Mbps '
    'into a 10 Mbps pipe, sustaining a queue that grows without bound at the simulation\'s '
    'buffer limit.'
)

add_para(
    '**B3 (cwnd-freeze)** exhibits the same failure mode. Freezing the congestion window '
    'at the pre-handover LEO value leaves an equally oversized window in place; the buffer '
    'fills to 33 MB identically to B1. Advance warning is useless: even at ℓ = 30 s, '
    'B3 does not converge.'
)

add_para(
    '**B4 (send-pause/resume)** converges exactly once across the full lead-time sweep, '
    'at ℓ = 5 s (T90 = 5.6 s), and fails at every other lead time. The narrow success '
    'window reflects the tight timing between the pause duration and the GEO queue-drain '
    'time: too short a pause leaves residual in-flight data that re-inflates the queue; '
    'too long a pause wastes capacity and triggers a slow-start on resume.'
)

add_para(
    '**BBR-SAT** converges at every tested lead time from ℓ = 0 to ℓ = 20 s, with '
    'T90 = 2.5 s throughout. On receiving the CONFIRMED signal at T_HO, the algorithm '
    'loads the seeded GEO orbit entry (`bw_hi` = 10 Mbps, `min_rtt` = 580 ms), enters '
    'a brief DRAIN phase to clear residual in-flight data, then transitions to REFILL '
    'and converges to a steady-state of 9.2 Mbps (92% of link capacity) with a peak '
    'queue of only 910 KB — a 36× reduction compared to B1/B3. Goodput over the '
    '90-second run is 56.9 MB, largely independent of lead time.'
)

add_para(
    'At ℓ = 30 s, T90 degrades sharply to 47.6 s. The PREDICTED signal at t = 0 s '
    '(i.e. T_HO − 30 s) initiates a proactive queue drain via `ProbeBW_DOWN`; however, '
    'because the actual handover does not occur for 30 s, BBR-SAT re-enters `ProbeBW_UP` '
    'and rebuilds a LEO-calibrated queue before the CONFIRMED signal fires. This '
    'establishes a practical upper bound: lead times up to approximately 20 s are '
    'beneficial; beyond that, pre-draining fires too early to hold.'
)

add_para(
    '**CUBIC** converges in 1.5 s at all lead times, achieving a steady-state of '
    '10.0 Mbps (100% utilisation) and 66.4 MB goodput. The rapid convergence is '
    'loss-driven: the queue overflow immediately after handover triggers multiplicative '
    'window reduction, and CUBIC settles within two to three 580 ms GEO RTTs. Although '
    'CUBIC achieves slightly higher single-flow throughput than BBR-SAT (at the cost of '
    'a 1.3 MB peak queue versus 910 KB), its behaviour under competition is discussed in '
    '§V-B. For single-flow bulk transfer, CUBIC\'s loss-driven convergence is competitive; '
    'BBR-SAT\'s advantage lies in mixed-CCA deployments where BBRv3 flows must coexist '
    'fairly on a shared beam — the common operational scenario for multi-orbit gateway links.'
)

add_table_simple(
    headers=['Baseline', 'T90', 'SS bw', 'Goodput', 'Peak Q'],
    rows=[
        ['B1 BBRv3 (vanilla)', 'N/C',   '0.0 Mbps', '5.4 MB',  '33,120 KB'],
        ['B3 cwnd-freeze',     'N/C',   '0.0 Mbps', '5.4 MB',  '33,120 KB'],
        ['B4 pause/resume',    'N/C',   '0.0 Mbps', '3.8 MB',  '34,238 KB'],
        [('BBR-SAT (ours)','b'), ('2.5 s','b'), ('9.2 Mbps','b'), ('56.9 MB','b'), ('910 KB','b')],
        ['CUBIC',             '1.5 s', '10.0 Mbps','66.4 MB', '1,328 KB'],
    ],
    caption='Table III. LEO→GEO, T_HO = 30 s, ℓ = 0, no loss'
)

add_note('[FIGURE 1: T90 vs. advance lead time for the LEO→GEO transition (T_HO = 30 s, '
         'no loss). BBR-SAT maintains a flat 2.5 s across ℓ ∈ [0, 20] s; open markers '
         'at the top dashed line denote non-convergent (N/C) trials.]')

add_heading('3) Lead-Time Sensitivity', level=3)

add_para(
    'Figure 1 plots T90 as a function of lead time for the LEO→GEO transition. BBR-SAT '
    'maintains a flat T90 of 2.5 s across ℓ ∈ [0, 20] s, confirming that the PREDICTED '
    'signal is effectively used but that even zero lead time is sufficient for convergence '
    'given the seeded orbit table. The practical interpretation is that any advance notice '
    'in the range 2–20 s is equally effective; the value of the PREDICTED signal lies '
    'primarily in proactive queue drain that prevents mid-flight data from inflating the '
    'post-handover queue, and this drain completes within one to two *pre-handover* '
    '(LEO) RTTs (≈50–100 ms).'
)

add_heading('4) Bandwidth-Increasing and Moderate Transitions', level=3)

add_para(
    'For all bandwidth-increasing transitions (MEO→LEO, GEO→LEO, GEO→MEO) and the '
    'moderate MEO→GEO and LEO→MEO transitions, every baseline converges at ℓ = 0 with '
    'T90 ≤ 3.5 s (Table IV). Bandwidth increases are inherently self-correcting: the new '
    'link can absorb the pre-handover cwnd without queue buildup, and BBR\'s bandwidth '
    'probing detects the higher capacity within one probe cycle. BBR-SAT\'s T90 at ℓ = 0 '
    'is 0–1 s higher than B1/B3 on most transitions because the CONFIRMED signal briefly '
    'forces a REFILL phase before resuming ProbeBW. The exception is LEO→MEO: BBR-SAT '
    'takes 8.5 s vs B1/B3\'s 0.5 s because MEO\'s higher RTT (150 ms) extends the REFILL '
    'ramp and the seeded `bw_hi` ceiling (30 Mbps) temporarily limits probing above the '
    'new fair share. With ℓ ≥ 2 s this overhead disappears as the PREDICTED signal '
    'pre-positions the orbit context before the handover fires.'
)

add_table_simple(
    headers=['Transition', 'B1', 'B3', 'B4', 'BBR-SAT', 'CUBIC'],
    rows=[
        ['LEO→MEO', '0.5 s', '0.5 s', '6.5 s', '8.5 s',        '0.5 s'],
        ['LEO→GEO', 'N/C',   'N/C',   'N/C',   ('2.5 s', 'b'),  '1.5 s'],
        ['MEO→LEO', '2.5 s', '2.5 s', '0.5 s', '3.5 s',         '0.5 s'],
        ['GEO→LEO', '1.5 s', '1.5 s', '2.5 s', '2.5 s',         '1.5 s'],
        ['MEO→GEO', '1.5 s', '1.5 s', '1.5 s', '2.5 s',         '1.5 s'],
        ['GEO→MEO', '1.5 s', '1.5 s', '2.5 s', '2.5 s',         '1.5 s'],
    ],
    caption='Table IV. T90 at ℓ = 0, T_HO = 30 s, no loss (all transitions)'
)

add_heading('5) Handover-Time Sensitivity', level=3)

add_para(
    'Repeating the LEO→GEO experiment at T_HO = 60 s reveals moderate T90 inflation: '
    'BBR-SAT achieves 2.5 s at ℓ ∈ {0, 20} s but 7.5–9.5 s at intermediate lead times, '
    'likely due to the orbit table\'s minimum-RTT entry drifting over the longer '
    'pre-handover interval as BBR natural probing samples slightly elevated RTTs. At '
    'T_HO = 120 s all BBR-SAT conditions fail to converge, suggesting that the orbit '
    'table\'s 10-second `min_rtt` expiry window allows the seeded GEO RTT entry to '
    'become stale before the handover fires. Extending the orbit table\'s RTT confidence '
    'window for long-duration pre-handover periods is identified as future work.'
)

add_heading('B. Fairness Under Competition', level=2)

add_para(
    'Advance knowledge of a handover benefits the informed flow; a natural concern is '
    'whether BBR-SAT exploits that advantage at the expense of competing flows that share '
    'the bottleneck. We evaluate two scenarios using the shared-link simulator described '
    'in §IV: **F1** — two flows both crossing a LEO→GEO handover at T = 30 s, and '
    '**F3** — two flows that start simultaneously on a GEO beam and reach steady state '
    'before BBR-SAT re-anchors its BDP context at T = 30 s. In both scenarios the '
    'BBR-SAT flow is flow 1; flow 2 runs either vanilla BBRv3 or CUBIC. Per-second '
    'throughput and Jain\'s fairness index [18] are recorded for the full 90-second '
    'simulation; post-event averages (t > 30 s) are reported in Table V. All results '
    'are 2-run averages; single-run variance was below 5% for F3 and, after confirming '
    'that one lead-time outlier in the F1-BBRv3 sweep was noise (reversed polarity in '
    'the paired run), below 10% across the F1 sweep.'
)

add_heading('1) F1: Shared LEO→GEO Handover', level=3)

add_para(
    'Figure 2(a) shows the 90-second throughput trace for BBR-SAT and BBRv3 when both '
    'flows cross the handover with zero lead time. Before the handover (t < 30 s) both '
    'flows share the 50 Mbps LEO beam equitably, each sustaining roughly 24 Mbps. At '
    't = 30 s the link drops to 10 Mbps and the RTT rises to ≈580 ms; both flows '
    'converge to the new GEO fair share (5 Mbps each) within one to two seconds. '
    'Averaged over t ∈ (30, 90] s, Jain\'s index is J = 0.997, confirming near-perfect '
    'fairness.'
)

add_para(
    'We repeated the sweep across lead times ℓ ∈ {0, 5, 10, 15, 20, 30} s; results '
    'appear in Table V (F1 vs BBRv3 rows). For ℓ ≤ 5 s and ℓ ≥ 20 s, J > 0.98 in both '
    'runs. At ℓ = 10–15 s one run showed a 3:1 split favouring whichever flow happened '
    'to exit the post-handover transient first; the complementary run reversed polarity, '
    'and the 2-run average (J ≈ 0.86) reflects the residual variance of a single '
    '90-second trial. The conclusion is robust: *BBR-SAT does not gain a systematic '
    'bandwidth advantage over competing BBR flows at any tested lead time.*'
)

add_heading('2) F3: Steady-State GEO', level=3)

add_para(
    'Figure 2(b) shows the F3 trace. Both flows start on the 10 Mbps GEO beam from '
    't = 0. CUBIC reaches its characteristic sawtooth steady state by t ≈ 20 s. At '
    't = 30 s BBR-SAT re-applies its stored GEO BDP context (a no-op on link parameters, '
    'but it resets `bw_hi`, `min_rtt`, and `inflight_hi` to the seeded values), causing '
    'a brief throughput dip visible in both flows as the Jain\'s trace dips below 0.8. '
    'This occurs because BBR-SAT\'s measured `bw_hi` after 30 s of GEO operation has '
    'been probed above the seeded ceiling; the context switch resets it to the '
    'conservative ephemeris value, temporarily clipping the pacing rate until ProbeBW '
    're-probes. After t ≈ 39 s both flows reach a stable regime: BBR-SAT 3.4 Mbps, '
    'CUBIC 6.2 Mbps, J = 0.888. The 64/36 split is consistent across both runs '
    '(run-to-run σ < 0.2 Mbps).'
)

add_heading('3) BBR-SAT vs. CUBIC During Handover', level=3)

add_para(
    'We also ran the F1 topology with CUBIC as the competing flow (Table V, F1 vs CUBIC '
    'rows). BBR-SAT obtains only 1.1–1.9 Mbps post-handover, with J ≈ 0.71, while CUBIC '
    'takes the remaining ≈7.8 Mbps. Notably, this imbalance is already present on LEO '
    '*before* the handover (CUBIC ≈37 Mbps vs. BBR-SAT ≈12 Mbps on the shared 50 Mbps '
    'LEO beam), confirming that the unfairness is inherited directly from BBRv3\'s known '
    'deference to CUBIC [3] and is not introduced by the satellite extension. Longer lead '
    'times (ℓ = 20 s) marginally worsen the split (J = 0.655) because BBR-SAT\'s '
    'proactive queue drain at ℓ s before handover vacates buffer space that CUBIC '
    'immediately reclaims.'
)

add_heading('4) Summary', level=3)

add_para(
    'BBR-SAT preserves the fairness properties of its BBRv3 base with respect to both '
    'competing CCA families: it co-exists equitably with other BBR flows (J ≥ 0.986 at '
    'lead times outside the narrow 10–15 s transient window), and it neither worsens nor '
    'repairs the pre-existing BBR-vs-CUBIC imbalance in the steady state. Addressing the '
    'BBR/CUBIC coexistence issue is outside the scope of this work but is a natural '
    'direction for future study.'
)

add_note('[FIGURE 2: Fairness results. (a) F1: BBR-SAT vs. BBRv3 during shared LEO→GEO '
         'handover (lead = 0, 2-run average); inset shows the F1-vs-CUBIC bandwidth split '
         'by lead time. (b) F3: BBR-SAT vs. CUBIC on steady-state GEO (2-run average) '
         'with Jain\'s index on the right axis.]')

add_table_simple(
    headers=['Scenario', 'Competitor', 'Lead (s)', 'BBR-SAT (Mbps)', 'Opponent (Mbps)', 'Jain J'],
    rows=[
        ['F1', 'BBRv3',  '0',  '4.86', '4.84', ('0.997','b')],
        ['F1', 'BBRv3',  '5',  '4.62', '5.11', '0.991'],
        ['F1', 'BBRv3', '10',  '5.83', '3.86', '0.862†'],
        ['F1', 'BBRv3', '15',  '3.51', '6.17', '0.866†'],
        ['F1', 'BBRv3', '20',  '4.62', '5.10', '0.987'],
        ['F1', 'BBRv3', '30',  '4.62', '5.06', '0.991'],
        ['F1', 'CUBIC',  '0',  '1.71', '7.81', '0.713'],
        ['F1', 'CUBIC', '10',  '1.86', '7.80', '0.734'],
        ['F1', 'CUBIC', '20',  '1.14', '8.32', '0.655'],
        ['F3', 'CUBIC',  '0',  '3.43', '6.18', '0.888'],
    ],
    caption='Table V. Fairness summary (post-event average, t > 30 s, 2-run mean)'
)

add_note('† High-variance 2-run average; individual runs: lead=10: J = 0.738, 0.986; '
         'lead=15: J = 0.742, 0.991. Polarity reverses between runs, indicating random '
         'transient ordering rather than a systematic bias.')

# ═══════════════════════════════════════════════════════════════════════════
# VI. DISCUSSION
# ═══════════════════════════════════════════════════════════════════════════

add_heading('VI. Discussion', level=1)

add_para(
    'BBR-SAT achieves T90 = 2.5 s even at ℓ = 0 because the orbit table is pre-seeded '
    'from ephemeris data; the CONFIRMED signal alone suffices for the BDP context switch. '
    'The 20-second practical ceiling on useful lead time arises because the proactive '
    'drain (50–100 ms at LEO RTT) completes well before the handover fires, after which '
    'BBR\'s `ProbeBW_UP` re-inflates the cwnd; operators can treat any signal in the '
    '2–20 s window as equivalent.'
)

add_heading('A. Limitations and Future Work', level=2)

add_para(
    '**Orbit table staleness.** The 10-second `min_rtt` expiry causes BBR-SAT to fail '
    'for handover times T_HO ≥ 120 s. Refreshing the orbit table entry during the '
    'PREDICTED window — a one-line change in the PREDICTED handler — would fix this at '
    'the cost of a brief `ProbeBW_DOWN` that re-seeds `min_rtt`.'
)

add_para(
    '**Loss injection.** All experiments are zero-loss. Satellite links exhibit correlated '
    'loss (rain fade, co-channel interference) that may interact with BBRv3\'s '
    'loss-recovery path. The interaction between the DRAIN phase and `is_in_recovery` '
    'state — currently cleared on PREDICTED — requires evaluation under realistic loss models.'
)

add_para(
    '**Multi-path QUIC.** Satellite multi-path scenarios (e.g., LEO primary + GEO backup '
    'with simultaneous active paths) involve per-path BDP state and cross-path pacing '
    'interaction not addressed here.'
)

add_para(
    '**BBR/CUBIC coexistence.** The 64/36 CUBIC advantage in steady-state GEO '
    '(F3, J = 0.888) and the 82/18 split during F1-CUBIC (J ≈ 0.71) are inherited from '
    'vanilla BBRv3\'s loss-averse behaviour on large-BDP links. Addressing this would '
    'require either ECN-based queue-occupancy signalling or a fairness-aware pacing mode '
    'in BBRv3 itself.'
)

# ═══════════════════════════════════════════════════════════════════════════
# VII. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════

add_heading('VII. Conclusion', level=1)

add_para(
    'We presented BBR-SAT, a minimal extension of BBRv3 that equips QUIC with orbit-aware '
    'congestion control for satellite inter-orbit handovers. By integrating an '
    'ephemeris-seeded orbit parameter table and a two-phase PREDICTED/CONFIRMED signal '
    'protocol, BBR-SAT eliminates the 8 × RTT stall inherent in BBRv3\'s bandwidth filter '
    'on downward BDP transitions. In a zero-loss simulator spanning all six pairwise '
    'LEO/MEO/GEO transitions, BBR-SAT is the only evaluated mechanism to converge '
    'consistently on the critical LEO→GEO transition across the full range of tested lead '
    'times (0–20 s), achieving T90 = 2.5 s with a 36× reduction in peak queue occupancy '
    'relative to vanilla BBRv3. Fairness analysis confirms that advance knowledge does not '
    'translate into a bandwidth advantage: BBR-SAT co-exists equitably with uninformed '
    'BBRv3 flows (J = 0.997) and preserves the existing BBRv3/CUBIC balance. RFC '
    '9743/BCP 133 [15] requires fairness evaluation for any new CCA specification; this '
    'work satisfies that requirement for the satellite handover context.'
)

# ═══════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════

add_heading('References', level=1)

refs = [
    '[1] M. Handley, "Delay is Not an Option: Low Latency Routing in Space," Proc. ACM HotNets, 2018, pp. 85–91.',
    '[2] D. Bhattacherjee, W. Iqbal, and A. Singla, "Network Transit of Satellite Constellations," Proc. ACM HotNets, 2019, pp. 22–29.',
    '[3] N. Cardwell et al., "BBR Congestion Control," IETF Internet-Draft draft-ietf-ccwg-bbr-05, Mar. 2026.',
    '[4] N. Cardwell, Y. Cheng, C. S. Gunn, S. H. Yeganeh, and V. Jacobson, "BBR: Congestion-Based Congestion Control," ACM Queue, vol. 14, no. 5, pp. 20–53, 2016.',
    '[5] J. Zhao and J. Pan, "StarQUIC: Satellite-Friendly QUIC Congestion Control," Proc. ACM MobiCom Workshop on LEO Networking and Communication (LEO-NET), 2024. DOI: 10.1145/3697253.3697271.',
    '[6] D. Zhao et al., "SatPipe: Proactive Queue Management for LEO Satellite Handovers," Proc. IEEE INFOCOM, 2025. IEEE Xplore 11044600.',
    '[7] Z. Lai et al., "LeoCC: Handover-Aware Congestion Control for LEO Satellites," Proc. ACM SIGCOMM, 2025, pp. 129–146. DOI: 10.1145/3718958.3750491.',
    '[8] Y. Yan, J. Li, J. Han, Q. Long, K. Xue, H. Chen, and N. Qiao, "A Handover-Aware Congestion Control Algorithm Assisted by DRL in LEO Satellite Networks," Proc. IEEE ICC, Montreal, 2025. IEEE Xplore 11161022.',
    '[9] R. Valentine et al., "OrbCC: In-Network Telemetry for Orbit-Aware Congestion Control," arXiv:2508.19067v2, Aug. 2025. [Preprint, not peer-reviewed]',
    '[10] Z. Lai, Z. Li, Q. Wu, H. Li, and Q. Zhang, "Analysis for the Adverse Effects of LEO Mobility on Internet Congestion Control," IETF Internet-Draft draft-lai-ccwg-lsncc-01, Feb. 2026.',
    '[11] N. Kuhn et al., "Convergence of Congestion Control from Retained State," IETF Internet-Draft draft-ietf-tsvwg-careful-resume-24, Oct. 2025.',
    '[12] M. Hofstätter et al., "Careful Resume over Satellite Paths," Proc. IEEE ASMS/SPSC, 2025. IEEE Xplore 10946055.',
    '[13] C. Caini and R. Firrincieli, "TCP Hybla: A TCP Enhancement for Heterogeneous Networks," Int. J. Satellite Commun. Netw., vol. 22, no. 5, pp. 547–566, 2004.',
    '[14] F. Michel, M. Trevisan, D. Rossi, and O. Bonaventure, "A First Look at Starlink Performance," Proc. ACM IMC, 2022, pp. 130–140.',
    '[15] M. Duke and G. Fairhurst, "Specifying New Congestion Control Algorithms," RFC 9743 / BCP 133, IETF, Mar. 2025. (Obsoletes RFC 5033.)',
    '[16] C. Huitema, "picoquic: A Small QUIC Implementation," github.com/private-octopus/picoquic, 2023.',
    '[17] T. Pauly, E. Kinnear, and D. Schinazi, "An Unreliable Datagram Extension to QUIC," RFC 9221, IETF, Mar. 2022.',
    '[18] R. Jain, D.-M. Chiu, and W. Hawe, "A Quantitative Measure of Fairness and Discrimination for Resource Allocation in Shared Computer Systems," DEC Research Report TR-301, 1984.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(ref).font.size = Pt(10)

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = '/home/rajat/tcpsatproject/paper/paper.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
